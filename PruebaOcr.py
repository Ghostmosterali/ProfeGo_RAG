"""
Módulo de procesamiento OCR con soporte asíncrono para archivos grandes
"""

import cv2
import pytesseract
import os
import docx
import PyPDF2
import pandas as pd
from pathlib import Path
import xml.etree.ElementTree as ET
import json
import asyncio
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from typing import Dict, Optional
import time
import logging

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuración de procesamiento
LARGE_FILE_THRESHOLD = 5 * 1024 * 1024  # 5MB - archivos considerados "grandes"
MAX_WORKERS = 4  # Número máximo de workers para procesamiento paralelo

class DocumentConverter:
    """Módulo de conversión de documentos con soporte asíncrono"""
    
    def __init__(self):
        # Configurar Tesseract si es necesario
        # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        # Extensiones soportadas
        self.image_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.webp'}
        self.document_extensions = {'.pdf', '.docx', '.doc', '.txt', '.csv', '.xlsx', '.xls', '.json', '.xml', '.odt'}
        
        # Executor para procesamiento asíncrono
        self.thread_executor = ThreadPoolExecutor(max_workers=MAX_WORKERS)
        self.process_executor = ProcessPoolExecutor(max_workers=2)
    
    def detect_file_type(self, file_path):
        """Detecta el tipo de archivo basado en su extensión"""
        extension = Path(file_path).suffix.lower()
        
        if extension in self.image_extensions:
            return 'image'
        elif extension == '.pdf':
            return 'pdf'
        elif extension in {'.docx', '.doc'}:
            return 'word'
        elif extension == '.txt':
            return 'text'
        elif extension in {'.csv'}:
            return 'csv'
        elif extension in {'.xlsx', '.xls'}:
            return 'excel'
        elif extension == '.json':
            return 'json'
        elif extension == '.xml':
            return 'xml'
        elif extension == '.odt':
            return 'odt'
        else:
            return 'unknown'
    
    def is_large_file(self, file_path: str) -> bool:
        """Determina si un archivo es grande y necesita procesamiento asíncrono"""
        try:
            file_size = os.path.getsize(file_path)
            return file_size > LARGE_FILE_THRESHOLD
        except:
            return False
    
    def extract_text_from_image(self, image_path):
        """Extrae texto de imágenes usando OCR"""
        try:
            img = cv2.imread(image_path)
            if img is None:
                raise ValueError(f"No se pudo cargar la imagen: {image_path}")
            
            # Convertir a escala de grises
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Aplicar umbral para mejorar el OCR
            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
            
            # Realizar OCR
            try:
                text = pytesseract.image_to_string(thresh, lang='spa')
            except pytesseract.TesseractError:
                try:
                    text = pytesseract.image_to_string(thresh, lang='eng')
                except pytesseract.TesseractError:
                    text = pytesseract.image_to_string(thresh)
            
            # Normalizar encoding
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text.strip()
        except Exception as e:
            return f"Error al procesar imagen: {str(e)}"
    
    async def extract_text_from_pdf_async(self, pdf_path):
        """Extrae texto de PDFs grandes de forma asíncrona"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.thread_executor,
            self._extract_text_from_pdf_sync,
            pdf_path
        )
    
    def _extract_text_from_pdf_sync(self, pdf_path):
        """Versión síncrona de extracción de PDF"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"Procesando PDF con {total_pages} páginas...")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Página {page_num + 1} ---\n"
                            text += page_text + "\n"
                        
                        # Log de progreso cada 10 páginas
                        if (page_num + 1) % 10 == 0:
                            logger.info(f"Procesadas {page_num + 1}/{total_pages} páginas")
                            
                    except Exception as e:
                        text += f"\n--- Error en página {page_num + 1}: {str(e)} ---\n"
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text.strip() if text.strip() else "No se pudo extraer texto del PDF"
        except Exception as e:
            return f"Error al procesar PDF: {str(e)}"
    
    def extract_text_from_pdf(self, pdf_path):
        """Extrae texto de archivos PDF (versión síncrona para PDFs pequeños)"""
        return self._extract_text_from_pdf_sync(pdf_path)
    
    def extract_text_from_word(self, word_path):
        """Extrae texto de documentos Word (.docx)"""
        try:
            doc = docx.Document(word_path)
            text = ""
            
            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extraer texto de tablas
            for table in doc.tables:
                text += "\n--- Tabla ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text.strip()
        except Exception as e:
            return f"Error al procesar documento Word: {str(e)}"
    
    def extract_text_from_csv(self, csv_path):
        """Extrae texto de archivos CSV"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(csv_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                df = pd.read_csv(csv_path, encoding='utf-8', errors='ignore')
            
            text = f"Archivo CSV: {os.path.basename(csv_path)}\n"
            text += f"Filas: {len(df)}, Columnas: {len(df.columns)}\n\n"
            text += "--- Columnas ---\n"
            text += ", ".join(df.columns.tolist()) + "\n\n"
            text += "--- Datos ---\n"
            text += df.to_string(index=False)
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text
        except Exception as e:
            return f"Error al procesar CSV: {str(e)}"
    
    def extract_text_from_excel(self, excel_path):
        """Extrae texto de archivos Excel"""
        try:
            excel_file = pd.ExcelFile(excel_path)
            text = f"Archivo Excel: {os.path.basename(excel_path)}\n"
            text += f"Hojas: {len(excel_file.sheet_names)}\n\n"
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                text += f"=== Hoja: {sheet_name} ===\n"
                text += f"Filas: {len(df)}, Columnas: {len(df.columns)}\n"
                text += "Columnas: " + ", ".join(df.columns.tolist()) + "\n\n"
                text += df.to_string(index=False) + "\n\n"
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text
        except Exception as e:
            return f"Error al procesar Excel: {str(e)}"
    
    def extract_text_from_json(self, json_path):
        """Extrae texto de archivos JSON"""
        try:
            with open(json_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            text = f"Archivo JSON: {os.path.basename(json_path)}\n\n"
            text += json.dumps(data, indent=2, ensure_ascii=False)
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text
        except Exception as e:
            return f"Error al procesar JSON: {str(e)}"
    
    def extract_text_from_xml(self, xml_path):
        """Extrae texto de archivos XML"""
        try:
            tree = ET.parse(xml_path)
            root = tree.getroot()
            
            text = f"Archivo XML: {os.path.basename(xml_path)}\n"
            text += f"Elemento raíz: {root.tag}\n\n"
            
            def extract_text_recursive(element, level=0):
                result = ""
                indent = "  " * level
                
                result += f"{indent}<{element.tag}>\n"
                
                if element.text and element.text.strip():
                    result += f"{indent}  {element.text.strip()}\n"
                
                for child in element:
                    result += extract_text_recursive(child, level + 1)
                
                return result
            
            text += extract_text_recursive(root)
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text
        except Exception as e:
            return f"Error al procesar XML: {str(e)}"
    
    def extract_text_from_text(self, text_path):
        """Lee archivos de texto plano"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text = None
            
            for encoding in encodings:
                try:
                    with open(text_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                with open(text_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
            
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            return text
        except Exception as e:
            return f"Error al procesar archivo de texto: {str(e)}"

# ===============================
# FUNCIONES PRINCIPALES
# ===============================

async def process_file_to_txt_async(file_path, output_path=None):
    """
    Versión asíncrona para archivos grandes
    """
    converter = DocumentConverter()
    
    response = {
        'success': False,
        'output_file': None,
        'extracted_text': '',
        'file_type': None,
        'error': None,
        'processing_time': 0
    }
    
    start_time = time.time()
    
    try:
        if not os.path.exists(file_path):
            response['error'] = f"El archivo '{file_path}' no existe"
            return response
        
        file_type = converter.detect_file_type(file_path)
        response['file_type'] = file_type
        
        if file_type == 'unknown':
            response['error'] = f"Tipo de archivo no soportado: {Path(file_path).suffix}"
            return response
        
        if output_path is None:
            base_name = Path(file_path).stem
            timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"{base_name}_converted_{timestamp}.txt"
        
        # Procesar según tipo y tamaño
        is_large = converter.is_large_file(file_path)
        logger.info(f"Procesando archivo {'grande' if is_large else 'normal'}: {file_path}")
        
        # Procesamiento asíncrono para PDFs grandes
        if file_type == 'pdf' and is_large:
            extracted_text = await converter.extract_text_from_pdf_async(file_path)
        else:
            # Procesamiento síncrono para el resto
            if file_type == 'image':
                extracted_text = converter.extract_text_from_image(file_path)
            elif file_type == 'pdf':
                extracted_text = converter.extract_text_from_pdf(file_path)
            elif file_type == 'word':
                extracted_text = converter.extract_text_from_word(file_path)
            elif file_type == 'text':
                extracted_text = converter.extract_text_from_text(file_path)
            elif file_type == 'csv':
                extracted_text = converter.extract_text_from_csv(file_path)
            elif file_type == 'excel':
                extracted_text = converter.extract_text_from_excel(file_path)
            elif file_type == 'json':
                extracted_text = converter.extract_text_from_json(file_path)
            elif file_type == 'xml':
                extracted_text = converter.extract_text_from_xml(file_path)
        
        if not extracted_text or extracted_text.strip() == "":
            response['error'] = f"No se pudo extraer texto del archivo {file_type}"
            return response
        
        response['extracted_text'] = extracted_text
        
        # Generar archivo .txt
        with open(output_path, 'w', encoding='utf-8-sig', errors='replace') as f:
            f.write(f"=== DOCUMENTO CONVERTIDO ===\n")
            f.write(f"Archivo original: {file_path}\n")
            f.write(f"Tipo de archivo: {file_type}\n")
            f.write(f"Fecha de conversión: {pd.Timestamp.now()}\n")
            f.write("="*50 + "\n\n")
            f.write(extracted_text)
        
        if os.path.exists(output_path):
            response['success'] = True
            response['output_file'] = output_path
            response['processing_time'] = time.time() - start_time
            logger.info(f"Procesamiento completado en {response['processing_time']:.2f} segundos")
        else:
            response['error'] = "Error al crear el archivo de salida"
            
    except Exception as e:
        response['error'] = f"Error inesperado durante el procesamiento: {str(e)}"
        response['processing_time'] = time.time() - start_time
    
    return response

def process_file_to_txt(file_path, output_path=None):
    """
    Versión síncrona mejorada (compatible con el código existente)
    """
    converter = DocumentConverter()
    
    # Si el archivo es grande, usar versión asíncrona
    if converter.is_large_file(file_path):
        logger.info(f"Archivo grande detectado, usando procesamiento optimizado")
        # Ejecutar versión asíncrona en un loop
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(process_file_to_txt_async(file_path, output_path))
        finally:
            loop.close()
    
    # Para archivos pequeños, usar procesamiento normal
    response = {
        'success': False,
        'output_file': None,
        'extracted_text': '',
        'file_type': None,
        'error': None
    }
    
    try:
        if not os.path.exists(file_path):
            response['error'] = f"El archivo '{file_path}' no existe"
            return response
        
        file_type = converter.detect_file_type(file_path)
        response['file_type'] = file_type
        
        if file_type == 'unknown':
            response['error'] = f"Tipo de archivo no soportado: {Path(file_path).suffix}"
            return response
        
        if output_path is None:
            base_name = Path(file_path).stem
            timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"{base_name}_converted_{timestamp}.txt"
        
        # Procesar archivo según su tipo
        extracted_text = ""
        
        if file_type == 'image':
            extracted_text = converter.extract_text_from_image(file_path)
        elif file_type == 'pdf':
            extracted_text = converter.extract_text_from_pdf(file_path)
        elif file_type == 'word':
            extracted_text = converter.extract_text_from_word(file_path)
        elif file_type == 'text':
            extracted_text = converter.extract_text_from_text(file_path)
        elif file_type == 'csv':
            extracted_text = converter.extract_text_from_csv(file_path)
        elif file_type == 'excel':
            extracted_text = converter.extract_text_from_excel(file_path)
        elif file_type == 'json':
            extracted_text = converter.extract_text_from_json(file_path)
        elif file_type == 'xml':
            extracted_text = converter.extract_text_from_xml(file_path)
        
        if not extracted_text or extracted_text.strip() == "":
            response['error'] = f"No se pudo extraer texto del archivo {file_type}"
            return response
        
        response['extracted_text'] = extracted_text
        
        # Generar archivo .txt
        with open(output_path, 'w', encoding='utf-8-sig', errors='replace') as f:
            f.write(f"=== DOCUMENTO CONVERTIDO ===\n")
            f.write(f"Archivo original: {file_path}\n")
            f.write(f"Tipo de archivo: {file_type}\n")
            f.write(f"Fecha de conversión: {pd.Timestamp.now()}\n")
            f.write("="*50 + "\n\n")
            f.write(extracted_text)
        
        if os.path.exists(output_path):
            response['success'] = True
            response['output_file'] = output_path
        else:
            response['error'] = "Error al crear el archivo de salida"
            
    except Exception as e:
        response['error'] = f"Error inesperado durante el procesamiento: {str(e)}"
    
    return response

def get_text_only(file_path):
    """
    Función que solo retorna el texto extraído (sin generar archivo)
    """
    converter = DocumentConverter()
    
    response = {
        'success': False,
        'text': '',
        'file_type': None,
        'error': None
    }
    
    try:
        if not os.path.exists(file_path):
            response['error'] = f"El archivo '{file_path}' no existe"
            return response
        
        file_type = converter.detect_file_type(file_path)
        response['file_type'] = file_type
        
        if file_type == 'unknown':
            response['error'] = f"Tipo de archivo no soportado: {Path(file_path).suffix}"
            return response
        
        # Extraer solo el texto
        if file_type == 'image':
            text = converter.extract_text_from_image(file_path)
        elif file_type == 'pdf':
            text = converter.extract_text_from_pdf(file_path)
        elif file_type == 'word':
            text = converter.extract_text_from_word(file_path)
        elif file_type == 'text':
            text = converter.extract_text_from_text(file_path)
        elif file_type == 'csv':
            text = converter.extract_text_from_csv(file_path)
        elif file_type == 'excel':
            text = converter.extract_text_from_excel(file_path)
        elif file_type == 'json':
            text = converter.extract_text_from_json(file_path)
        elif file_type == 'xml':
            text = converter.extract_text_from_xml(file_path)
        else:
            response['error'] = f"Tipo de archivo no implementado: {file_type}"
            return response
        
        response['text'] = text
        response['success'] = True
        
    except Exception as e:
        response['error'] = f"Error durante extracción: {str(e)}"
    
    return response

def check_supported_file(file_path):
    """
    Verifica si un archivo es soportado sin procesarlo
    """
    converter = DocumentConverter()
    
    extension = Path(file_path).suffix.lower()
    file_type = converter.detect_file_type(file_path)
    
    return {
        'supported': file_type != 'unknown',
        'file_type': file_type,
        'extension': extension
    }