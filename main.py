from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Header, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import pyrebase
import os
import re
from pathlib import Path
from dotenv import load_dotenv
from typing import List, Optional, Dict
import json
from datetime import datetime
import tempfile
import io
import time
import uuid
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rag_system import get_rag_system, initialize_rag_system
from typing import List, Dict, Optional
import json
from pathlib import Path
import firebase_admin
from firebase_admin import credentials, auth as admin_auth
# Importar el mÃ³dulo OCR
from PruebaOcr import process_file_to_txt, check_supported_file, get_text_only

# Importar el mÃ³dulo de Google Cloud Storage mejorado
from gcs_storage import GCSStorageManagerV2

# Importar el servicio de Gemini AI
from gemini_service import generar_plan_estudio

from email_service import (
    email_service,
    send_verification_email,
    send_password_reset_email,
    send_login_notification,
    verify_token,
    invalidate_token
)
from user_agents import parse as parse_user_agent
# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Cargar variables de entorno
load_dotenv()

firebase_creds_json = os.getenv("FIREBASE_ADMIN_CREDENTIALS_JSON")
if firebase_creds_json:
    cred = credentials.Certificate(json.loads(firebase_creds_json))
    firebase_admin.initialize_app(cred)
else:
    logger.error("âŒ FIREBASE_ADMIN_CREDENTIALS_JSON no configurada")

app = FastAPI(title="ProfeGo API", version="2.0.0")
rag_system = None

# Rate Limiter
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Configurar CORS para desarrollo y producciÃ³n
RENDER_EXTERNAL_URL = os.getenv("RENDER_EXTERNAL_URL")
allowed_origins = []

if RENDER_EXTERNAL_URL:
    allowed_origins = [
        RENDER_EXTERNAL_URL,
        f"https://{RENDER_EXTERNAL_URL.replace('https://', '')}",
    ]
else:
    allowed_origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ConfiguraciÃ³n de archivos
MAX_FILE_SIZE = 15 * 1024 * 1024  # 15MB
ALLOWED_EXTENSIONS = {
    '.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', 
    '.png', '.xlsx', '.xls', '.csv', '.json', '.xml'
}

# Obtener directorio base del script
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")

# Verificar que el directorio frontend existe
if not os.path.exists(FRONTEND_DIR):
    print(f"âš ï¸ ADVERTENCIA: No se encontrÃ³ el directorio frontend en {FRONTEND_DIR}")
else:
    print(f"âœ… Frontend encontrado en: {FRONTEND_DIR}")
    app.mount("/frontend", StaticFiles(directory=FRONTEND_DIR), name="frontend")
    print(f"âœ… Archivos estÃ¡ticos montados en /frontend")

# Firebase Config
firebaseConfig = {
    "apiKey": os.getenv("FIREBASE_API_KEY"),
    "authDomain": os.getenv("FIREBASE_AUTH_DOMAIN"),
    "projectId": os.getenv("FIREBASE_PROJECT_ID"),
    "storageBucket": os.getenv("FIREBASE_STORAGE_BUCKET"),
    "messagingSenderId": os.getenv("FIREBASE_MESSAGING_SENDER_ID"),
    "appId": os.getenv("FIREBASE_APP_ID"),
    "databaseURL": os.getenv("FIREBASE_DATABASE_URL", "")
}

firebase = pyrebase.initialize_app(firebaseConfig)
auth = firebase.auth()

# Google Cloud Storage Manager V2
gcs_storage = GCSStorageManagerV2(
    bucket_name=os.getenv("GCS_BUCKET_NAME", "bucket-profe-go")
)

# ---------------- Modelos Pydantic ----------------
class UserLogin(BaseModel):
    email: str
    password: str

class UserResponse(BaseModel):
    email: str
    token: str
    message: str

class FileInfo(BaseModel):
    name: str
    type: str
    size: str
    category: str
    date: str
    download_url: Optional[str] = None

class ProcessingResult(BaseModel):
    success: bool
    files_uploaded: int
    files_processed: int
    message: str
    errors: List[str] = []

class PaginatedFiles(BaseModel):
    files: List[FileInfo]
    total: int
    page: int
    pages: int
    per_page: int

class PlanGenerationRequest(BaseModel):
    plan_filename: str
    diagnostico_filename: Optional[str] = None

class PlanResponse(BaseModel):
    success: bool
    plan_id: Optional[str] = None
    plan_data: Optional[Dict] = None
    error: Optional[str] = None
    processing_time: Optional[float] = None

class PasswordResetRequest(BaseModel):
    email: str

class PasswordReset(BaseModel):
    token: str
    new_password: str

# ---------------- Utilidades ----------------
class ProfeGoUtils:
    @staticmethod
    def validar_email(email: str) -> bool:
        """Validar formato de email"""
        pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
        return re.match(pattern, email) is not None
    
    @staticmethod
    def validar_password(password: str) -> bool:
        """Validar que la contraseÃ±a tenga al menos 6 caracteres"""
        return len(password) >= 6
    
    @staticmethod
    def obtener_tipo_archivo(filename: str) -> str:
        """Determinar el tipo de archivo basado en su extensiÃ³n"""
        ext = os.path.splitext(filename)[1].lower()
        
        tipos = {
            '.pdf': "PDF",
            '.jpg': "Imagen", '.jpeg': "Imagen", '.png': "Imagen",
            '.doc': "Word", '.docx': "Word",
            '.xls': "Excel", '.xlsx': "Excel",
            '.txt': "Texto",
            '.csv': "CSV",
            '.json': "JSON",
            '.xml': "XML"
        }
        
        return tipos.get(ext, "Archivo")
    
    @staticmethod
    def validar_extension(filename: str) -> bool:
        """Validar si la extensiÃ³n del archivo estÃ¡ permitida"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in ALLOWED_EXTENSIONS

# ---------------- Dependency para autenticaciÃ³n ----------------
async def get_current_user(authorization: str = Header(None)):
    """Verificar token de Firebase y extraer usuario - REQUIERE EMAIL VERIFICADO"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Token no proporcionado")
    
    token = authorization.replace("Bearer ", "")
    
    try:
        # â­ USAR FIREBASE ADMIN SDK para obtener informaciÃ³n completa del token
        decoded_token = admin_auth.verify_id_token(token)
        email = decoded_token.get('email')
        email_verified = decoded_token.get('email_verified', False)
        
        # â­ VALIDACIÃ“N CRÃTICA: Verificar que el email estÃ© verificado
        if not email_verified:
            logger.warning(f"âš ï¸ Usuario {email} intentÃ³ acceder sin verificar su correo")
            raise HTTPException(
                status_code=403, 
                detail="Por favor verifica tu correo electrÃ³nico antes de acceder a la aplicaciÃ³n. Revisa tu bandeja de entrada."
            )
        
        logger.info(f"âœ… Usuario autenticado y verificado: {email}")
        return {"email": email, "token": token, "email_verified": email_verified}
        
    except admin_auth.InvalidIdTokenError:
        logger.error("Token de Firebase invÃ¡lido")
        raise HTTPException(status_code=401, detail="Token invÃ¡lido")
    except admin_auth.ExpiredIdTokenError:
        logger.error("Token de Firebase expirado")
        raise HTTPException(status_code=401, detail="Token expirado. Por favor inicia sesiÃ³n nuevamente")
    except HTTPException:
        # Re-lanzar excepciones HTTP (como el 403 de email no verificado)
        raise
    except Exception as e:
        logger.error(f"Error verificando token: {e}")
        raise HTTPException(status_code=401, detail="Error de autenticaciÃ³n")

# ============================================================================
# STARTUP EVENT - InicializaciÃ³n del sistema RAG
# ============================================================================

@app.on_event("startup")
async def startup_event():
    """
    Inicializa el sistema RAG al arrancar la aplicaciÃ³n
    """
    global rag_system, rag_analyzer
    
    logger.info("ðŸš€ Inicializando sistema RAG...")
    
    try:
        from pathlib import Path
        
        dirs_to_create = [
            './rag_data/cuentos',
            './rag_data/canciones',
            './rag_data/actividades',
            './rag_data/vector_db'
        ]
        
        for dir_path in dirs_to_create:
            Path(dir_path).mkdir(parents=True, exist_ok=True)
        
        logger.info("âœ… Directorios RAG creados/verificados")
        
        cuentos_count = len(list(Path('./rag_data/cuentos').glob('**/*.txt')))
        canciones_count = len(list(Path('./rag_data/canciones').glob('**/*.txt')))
        actividades_count = len(list(Path('./rag_data/actividades').glob('**/*.txt')))
        
        logger.info(f"ðŸ“š Biblioteca: {cuentos_count} cuentos, {canciones_count} canciones, {actividades_count} actividades")
        
        if cuentos_count == 0 and canciones_count == 0 and actividades_count == 0:
            logger.warning("âš ï¸ Biblioteca RAG vacÃ­a - No se encontraron archivos .txt")
            logger.warning("ðŸ’¡ Agrega archivos .txt en ./rag_data/cuentos, ./rag_data/canciones y ./rag_data/actividades")
            logger.warning("ðŸ’¡ Luego ejecuta: python init_rag.py")
        
        rag_system = initialize_rag_system()
        
        if rag_system is None:
            logger.warning("âš ï¸ Sistema RAG no pudo inicializarse")
            return
        
        rag_analyzer = RAGAnalyzer(rag_system)
        logger.info("âœ… RAG Analyzer inicializado")
        
        stats = rag_system.get_stats()
        
        if stats['total_documents'] == 0:
            logger.warning("âš ï¸ Vector store vacÃ­o - ejecuta 'python init_rag.py'")
        else:
            logger.info(f"âœ… Vector store listo: {stats['total_documents']} documentos indexados")
        
    except Exception as e:
        logger.error(f"âŒ Error inicializando RAG: {e}", exc_info=True)
        logger.warning("âš ï¸ La aplicaciÃ³n continuarÃ¡ sin RAG")
        rag_system = None
        rag_analyzer = None

# ============================================================================
# RUTAS DE AUTENTICACIÃ“N
# ============================================================================

@app.post("/api/auth/login", response_model=UserResponse)
@limiter.limit("10/minute")
async def login(request: Request, user_data: UserLogin):
    """Iniciar sesiÃ³n con notificaciÃ³n de email"""
    logger.info(f"ðŸ“§ Intento de login: {user_data.email}")
    
    if not ProfeGoUtils.validar_email(user_data.email):
        raise HTTPException(status_code=400, detail="Email invÃ¡lido")
    
    if not ProfeGoUtils.validar_password(user_data.password):
        raise HTTPException(status_code=400, detail="La contraseÃ±a debe tener al menos 6 caracteres")
    
    try:
        # Login en Firebase
        user = auth.sign_in_with_email_and_password(user_data.email, user_data.password)
        token = user.get("idToken")
        
        # â­ NUEVO: Verificar estado del email inmediatamente en el login
        try:
            decoded_token = admin_auth.verify_id_token(token)
            email_verified = decoded_token.get('email_verified', False)
            
            if not email_verified:
                logger.warning(f"âš ï¸ Usuario {user_data.email} intentÃ³ iniciar sesiÃ³n sin verificar email")
                raise HTTPException(
                    status_code=403,
                    detail="Por favor verifica tu correo electrÃ³nico antes de iniciar sesiÃ³n. Revisa tu bandeja de entrada."
                )
        except HTTPException:
            raise
        except Exception as verify_error:
            logger.error(f"Error verificando estado del email: {verify_error}")
            # Continuar con el login si hay error verificando (por compatibilidad)
        
        # Inicializar usuario en GCS
        gcs_storage.inicializar_usuario(user_data.email)
        logger.info(f"âœ… Login exitoso: {user_data.email}")
        
        # Enviar notificaciÃ³n de inicio de sesiÃ³n
        try:
            user_agent_string = request.headers.get('user-agent', '')
            client_ip = request.client.host if request.client else None
            
            try:
                from user_agents import parse
                ua = parse(user_agent_string)
                device_info = f"{ua.browser.family} {ua.browser.version_string} en {ua.os.family} {ua.os.version_string}"
            except:
                device_info = "Navegador desconocido"
            
            email_result = send_login_notification(
                email=user_data.email,
                device_info=device_info,
                ip_address=client_ip,
                location="Guadalajara, Jalisco, MX"
            )
            
            if email_result['success']:
                logger.info(f"ðŸ”” NotificaciÃ³n de login enviada a {user_data.email}")
        except Exception as e:
            logger.warning(f"âš ï¸ No se pudo enviar notificaciÃ³n de login: {e}")
        
        return UserResponse(
            email=user_data.email,
            token=token,
            message=f"Bienvenido/a {user_data.email}"
        )
    except Exception as e:
        error_msg = str(e)
        logger.error(f"âŒ Error en login: {error_msg}")
        
        if "EMAIL_NOT_FOUND" in error_msg:
            raise HTTPException(status_code=400, detail="Usuario no encontrado")
        elif "INVALID_PASSWORD" in error_msg or "INVALID_LOGIN_CREDENTIALS" in error_msg:
            raise HTTPException(status_code=400, detail="Credenciales incorrectas")
        elif "TOO_MANY_ATTEMPTS" in error_msg:
            raise HTTPException(status_code=429, detail="Demasiados intentos. Intenta mÃ¡s tarde")
        else:
            raise HTTPException(status_code=400, detail="Error de autenticaciÃ³n")

@app.post("/api/auth/register")
@limiter.limit("3/minute")
async def register(request: Request, user_data: UserLogin):
    """Registrar nuevo usuario con verificaciÃ³n de email"""
    logger.info(f"ðŸ“ Intento de registro: {user_data.email}")
    
    if not ProfeGoUtils.validar_email(user_data.email):
        raise HTTPException(status_code=400, detail="Email invÃ¡lido")
    
    if not ProfeGoUtils.validar_password(user_data.password):
        raise HTTPException(status_code=400, detail="La contraseÃ±a debe tener al menos 6 caracteres")
    
    try:
        # Registrar en Firebase
        auth.create_user_with_email_and_password(user_data.email, user_data.password)
        
        # Inicializar estructura en GCS
        gcs_storage.inicializar_usuario(user_data.email)
        
        logger.info(f"âœ… Usuario registrado: {user_data.email}")
        
        # â­ NUEVO: Enviar email de verificaciÃ³n
        try:
            email_result = send_verification_email(user_data.email)
            if email_result['success']:
                logger.info(f"ðŸ“§ Email de verificaciÃ³n enviado a {user_data.email}")
                return {
                    "message": "Usuario registrado correctamente. Por favor verifica tu correo electrÃ³nico para activar tu cuenta.",
                    "email_sent": True
                }
            else:
                logger.warning(f"âš ï¸ No se pudo enviar email de verificaciÃ³n: {email_result.get('error')}")
                return {
                    "message": "Usuario registrado correctamente, pero no se pudo enviar el email de verificaciÃ³n. Ya puedes iniciar sesiÃ³n.",
                    "email_sent": False
                }
        except Exception as e:
            logger.error(f"âŒ Error enviando email de verificaciÃ³n: {e}")
            return {
                "message": "Usuario registrado correctamente, pero hubo un problema con el email. Ya puedes iniciar sesiÃ³n.",
                "email_sent": False
            }
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"âŒ Error en registro: {error_msg}")
        
        if "EMAIL_EXISTS" in error_msg:
            raise HTTPException(status_code=400, detail="Este email ya estÃ¡ registrado")
        else:
            raise HTTPException(status_code=400, detail="Error en el registro")


@app.get("/api/auth/verify-email")
async def verify_email(token: str):
    """
    Verifica el email del usuario mediante token
    Esta ruta se llama cuando el usuario hace clic en el enlace del email
    """
    logger.info(f"ðŸ“§ Verificando email con token: {token[:10]}...")
    
    try:
        # Verificar token
        email = verify_token(token, 'verification')
        
        if not email:
            raise HTTPException(
                status_code=400,
                detail="Token invÃ¡lido o expirado. Por favor solicita un nuevo enlace de verificaciÃ³n."
            )
        
        # Marcar el email como verificado en Firebase Admin SDK
        try:
            # Buscar usuario por email
            user = admin_auth.get_user_by_email(email)
            
            # Actualizar el estado de verificaciÃ³n
            admin_auth.update_user(
                user.uid,
                email_verified=True
            )
            logger.info(f"âœ… Email marcado como verificado en Firebase: {email}")
        except Exception as firebase_error:
            logger.error(f"âŒ Error actualizando estado en Firebase: {firebase_error}")
            # Continuar aunque falle la actualizaciÃ³n en Firebase
        
        # Invalidar token despuÃ©s de usarlo
        invalidate_token(token)
        
        logger.info(f"âœ… Email verificado exitosamente: {email}")
        
        # â­ CORRECCIÃ“N: Usar ruta dinÃ¡mica del frontend
        verification_success_path = os.path.join(FRONTEND_DIR, "verification-success.html")
        
        if os.path.exists(verification_success_path):
            return FileResponse(
                verification_success_path,
                media_type="text/html"
            )
        else:
            # Si no existe la pÃ¡gina, devolver respuesta JSON como fallback
            return JSONResponse(
                content={
                    "success": True,
                    "message": "âœ… Email verificado exitosamente. Ya puedes iniciar sesiÃ³n.",
                    "email": email,
                    "redirect": "/login.html"
                },
                status_code=200
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error verificando email: {e}")
        raise HTTPException(status_code=500, detail="Error verificando email")

@app.post("/api/auth/request-password-reset")
@limiter.limit("3/minute")
async def request_password_reset(request: Request, reset_request: PasswordResetRequest):
    """
    Solicita restablecimiento de contraseÃ±a
    EnvÃ­a email con enlace para crear nueva contraseÃ±a
    """
    logger.info(f"ðŸ”‘ Solicitud de reset de contraseÃ±a para: {reset_request.email}")
    
    try:
        # Verificar que el email existe en Firebase
        # Nota: Firebase no tiene una API directa para verificar si un email existe
        # sin intentar un login, asÃ­ que enviamos el email sin verificar
        
        # Enviar email de restablecimiento
        email_result = send_password_reset_email(reset_request.email)
        
        if not email_result['success']:
            raise HTTPException(
                status_code=500,
                detail="Error enviando el correo de restablecimiento"
            )
        
        logger.info(f"âœ… Email de reset enviado a {reset_request.email}")
        
        # Por seguridad, siempre devolvemos el mismo mensaje
        # incluso si el email no existe (previene enumeraciÃ³n de usuarios)
        return {
            "message": "Si el correo existe en nuestro sistema, recibirÃ¡s un enlace para restablecer tu contraseÃ±a.",
            "email_sent": True
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error en solicitud de reset: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error procesando la solicitud"
        )

@app.post("/api/auth/reset-password")
@limiter.limit("5/minute")
async def reset_password(request: Request, reset_data: PasswordReset):
    """
    Restablece la contraseÃ±a del usuario usando Firebase Admin SDK
    """
    logger.info(f"ðŸ”‘ Procesando reset de contraseÃ±a...")
    
    try:
        # 1. Verificar el token personalizado
        email = verify_token(reset_data.token, 'reset_password')
        
        if not email:
            raise HTTPException(
                status_code=400,
                detail="Token invÃ¡lido o expirado. Solicita un nuevo enlace de restablecimiento."
            )
        
        logger.info(f"âœ… Token vÃ¡lido para: {email}")
        
        # 2. Validar la nueva contraseÃ±a
        if not ProfeGoUtils.validar_password(reset_data.new_password):
            raise HTTPException(
                status_code=400,
                detail="La contraseÃ±a debe tener al menos 6 caracteres"
            )
        
        # 3. Obtener el usuario por email usando Firebase Admin SDK
        try:
            user = admin_auth.get_user_by_email(email)  # â­ Cambiado aquÃ­
            logger.info(f"âœ… Usuario encontrado: {user.uid}")
        except Exception as e:
            logger.error(f"âŒ Usuario no encontrado: {e}")
            raise HTTPException(
                status_code=404,
                detail="Usuario no encontrado en el sistema"
            )
        
        # 4. Actualizar la contraseÃ±a usando Admin SDK
        try:
            admin_auth.update_user(  # â­ Cambiado aquÃ­
                user.uid,
                password=reset_data.new_password
            )
            logger.info(f"âœ… ContraseÃ±a actualizada correctamente para: {email}")
        except Exception as e:
            logger.error(f"âŒ Error actualizando contraseÃ±a: {e}")
            raise HTTPException(
                status_code=500,
                detail="Error actualizando la contraseÃ±a. Intenta nuevamente."
            )
        
        # 5. Invalidar el token para que no se pueda reutilizar
        invalidate_token(reset_data.token)
        logger.info(f"ðŸ—‘ï¸ Token invalidado")
        
        return {
            "success": True,
            "message": "ContraseÃ±a actualizada correctamente. Ya puedes iniciar sesiÃ³n con tu nueva contraseÃ±a."
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error inesperado al cambiar contraseÃ±a: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Error inesperado. Por favor intenta nuevamente."
        )

@app.post("/api/auth/resend-verification")
@limiter.limit("3/minute")
async def resend_verification(request: Request, user_data: UserLogin):
    """
    ReenvÃ­a el email de verificaciÃ³n
    """
    logger.info(f"ðŸ“§ Reenviando verificaciÃ³n a: {user_data.email}")
    
    try:
        # Enviar email de verificaciÃ³n
        email_result = send_verification_email(user_data.email)
        
        if not email_result['success']:
            raise HTTPException(
                status_code=500,
                detail="Error enviando el correo de verificaciÃ³n"
            )
        
        logger.info(f"âœ… Email de verificaciÃ³n reenviado a {user_data.email}")
        
        return {
            "message": "Email de verificaciÃ³n enviado. Por favor revisa tu bandeja de entrada.",
            "email_sent": True
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error reenviando verificaciÃ³n: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error enviando el correo"
        )
# ============================================================================
# RUTAS DE ARCHIVOS
# ============================================================================

@app.post("/api/files/upload", response_model=ProcessingResult)
@limiter.limit("10/minute")
async def upload_files(
    request: Request,
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Subir y procesar archivos con validaciones"""
    user_email = current_user["email"]
    
    archivos_subidos = []
    archivos_procesados = []
    errores_procesamiento = []
    
    for file in files:
        try:
            if not ProfeGoUtils.validar_extension(file.filename):
                errores_procesamiento.append(
                    f"{file.filename}: Tipo de archivo no permitido"
                )
                continue
            
            content = await file.read()
            
            if len(content) > MAX_FILE_SIZE:
                errores_procesamiento.append(
                    f"{file.filename}: Archivo muy grande (mÃ¡x: 15MB)"
                )
                continue
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp_file:
                tmp_file.write(content)
                tmp_file_path = tmp_file.name
            
            try:
                resultado_subida = gcs_storage.subir_archivo_desde_bytes(
                    contenido=content,
                    email=user_email,
                    nombre_archivo=file.filename,
                    es_procesado=False
                )
                
                if resultado_subida['success']:
                    archivos_subidos.append(file.filename)
                    
                    verificacion = check_supported_file(tmp_file_path)
                    
                    if verificacion['supported']:
                        nombre_base = Path(file.filename).stem
                        resultado_conversion = process_file_to_txt(tmp_file_path)
                        
                        if resultado_conversion['success']:
                            with open(resultado_conversion['output_file'], 'rb') as f:
                                contenido_procesado = f.read()
                            
                            resultado_txt = gcs_storage.subir_archivo_desde_bytes(
                                contenido=contenido_procesado,
                                email=user_email,
                                nombre_archivo=f"{nombre_base}_procesado.txt",
                                es_procesado=True
                            )
                            
                            if resultado_txt['success']:
                                archivos_procesados.append({
                                    'original': file.filename,
                                    'txt': f"{nombre_base}_procesado.txt"
                                })
                            
                            if os.path.exists(resultado_conversion['output_file']):
                                os.remove(resultado_conversion['output_file'])
                else:
                    errores_procesamiento.append(
                        f"{file.filename}: Error subiendo a GCS"
                    )
                    
            finally:
                if os.path.exists(tmp_file_path):
                    os.remove(tmp_file_path)
                    
        except Exception as ex:
            errores_procesamiento.append(f"{file.filename}: {str(ex)}")
    
    message = f"Archivos subidos: {len(archivos_subidos)}"
    if archivos_procesados:
        message += f", Procesados: {len(archivos_procesados)}"
    
    return ProcessingResult(
        success=len(archivos_subidos) > 0,
        files_uploaded=len(archivos_subidos),
        files_processed=len(archivos_procesados),
        message=message,
        errors=errores_procesamiento
    )

@app.get("/api/files/list")
async def list_files(
    page: int = Query(1, ge=1),
    per_page: int = Query(100, ge=1, le=100),
    current_user: dict = Depends(get_current_user)
):
    """Listar archivos del usuario"""
    user_email = current_user["email"]
    files_info = []
    
    try:
        archivos_originales = gcs_storage.listar_archivos(user_email, "uploads")
        archivos_procesados = gcs_storage.listar_archivos(user_email, "processed")
        
        for archivo in archivos_originales:
            files_info.append({
                "name": archivo['name'],
                "type": ProfeGoUtils.obtener_tipo_archivo(archivo['name']),
                "size": f"{archivo['size_mb']} MB",
                "category": "original",
                "date": archivo['date']
            })
        
        for archivo in archivos_procesados:
            if not archivo['name'].startswith('plan_') or not archivo['name'].endswith('.json'):
                files_info.append({
                    "name": archivo['name'],
                    "type": "TXT Procesado",
                    "size": f"{archivo['size_mb']} MB",
                    "category": "procesado",
                    "date": archivo['date']
                })
        
        return files_info
        
    except Exception as e:
        logger.error(f"âŒ Error listando archivos: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listando archivos: {str(e)}")

@app.get("/api/files/download/{category}/{filename}")
async def download_file(
    category: str,
    filename: str,
    current_user: dict = Depends(get_current_user)
):
    """Descargar archivo directamente desde GCS"""
    user_email = current_user["email"]
    
    try:
        es_procesado = category == "procesado"
        
        # Obtener el archivo desde GCS
        contenido = gcs_storage.obtener_archivo_bytes(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=es_procesado
        )
        
        if contenido is None:
            raise HTTPException(status_code=404, detail="Archivo no encontrado")
        
        # Determinar el tipo MIME
        content_type = "application/octet-stream"
        ext = Path(filename).suffix.lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.txt': 'text/plain',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xls': 'application/vnd.ms-excel',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }
        content_type = mime_types.get(ext, content_type)
        
        # Retornar el archivo como stream
        return StreamingResponse(
            io.BytesIO(contenido),
            media_type=content_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as ex:
        raise HTTPException(status_code=500, detail=f"Error descargando archivo: {str(ex)}")

@app.get("/api/files/preview/{category}/{filename}")
async def preview_file(
    category: str,
    filename: str,
    current_user: dict = Depends(get_current_user)
):
    """Vista previa de archivo - devuelve contenido segÃºn tipo"""
    user_email = current_user["email"]
    
    try:
        es_procesado = category == "procesado"
        
        # Obtener el archivo desde GCS
        contenido = gcs_storage.obtener_archivo_bytes(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=es_procesado
        )
        
        if contenido is None:
            raise HTTPException(status_code=404, detail="Archivo no encontrado")
        
        # Detectar tipo de archivo
        ext = Path(filename).suffix.lower()
        
        # Para PDFs e imÃ¡genes, devolver el archivo directamente
        if ext in ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            mime_types = {
                '.pdf': 'application/pdf',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.bmp': 'image/bmp'
            }
            
            return StreamingResponse(
                io.BytesIO(contenido),
                media_type=mime_types.get(ext, 'application/octet-stream'),
                headers={"Content-Disposition": f"inline; filename={filename}"}
            )
        
        # Para archivos TXT, devolver el contenido como JSON
        elif ext == '.txt':
            try:
                texto = contenido.decode('utf-8')
            except UnicodeDecodeError:
                texto = contenido.decode('latin-1', errors='ignore')
            
            return JSONResponse(content={
                "type": "text",
                "content": texto,
                "filename": filename
            })
        
        else:
            raise HTTPException(
                status_code=400, 
                detail="Tipo de archivo no soportado para vista previa"
            )
            
    except HTTPException:
        raise
    except Exception as ex:
        logger.error(f"Error en preview: {str(ex)}")
        raise HTTPException(status_code=500, detail=f"Error: {str(ex)}")

@app.delete("/api/files/delete/{category}/{filename}")
@limiter.limit("20/minute")
async def delete_file(
    request: Request,
    category: str,
    filename: str,
    current_user: dict = Depends(get_current_user)
):
    """Eliminar archivo de GCS"""
    user_email = current_user["email"]
    
    try:
        es_procesado = category == "procesado"
        
        resultado = gcs_storage.eliminar_archivo(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=es_procesado
        )
        
        if not resultado['success']:
            raise HTTPException(status_code=404, detail=resultado.get('error', 'Archivo no encontrado'))
        
        return {"message": f"Archivo '{filename}' eliminado correctamente"}
        
    except HTTPException:
        raise
    except Exception as ex:
        raise HTTPException(status_code=500, detail=f"Error eliminando archivo: {str(ex)}")
# ============================================================================
# RUTAS PARA GENERACIÃ“N DE PLANES CON IA + RAG
# ============================================================================

@app.post("/api/plans/generate", response_model=PlanResponse)
@limiter.limit("5/hour")
async def generate_plan_with_rag(
    request: Request,
    plan_file: UploadFile = File(..., description="Archivo del plan de estudios"),
    diagnostico_file: Optional[UploadFile] = File(None, description="Archivo de diagnÃ³stico (opcional)"),
    current_user: dict = Depends(get_current_user)
):
    """
    Genera un plan de estudio personalizado usando Gemini AI + RAG
    VERSIÃ“N CON SOPORTE PARA ACTIVIDADES
    """
    user_email = current_user["email"]
    start_time = time.time()
    
    logger.info(f"ðŸŽ“ Generando plan con RAG para usuario: {user_email}")
    
    try:
        # ========== VALIDACIÃ“N DE ARCHIVOS ==========
        
        if not ProfeGoUtils.validar_extension(plan_file.filename):
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de archivo no permitido para plan: {plan_file.filename}"
            )
        
        plan_content = await plan_file.read()
        if len(plan_content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail="El archivo del plan excede el lÃ­mite de 15MB"
            )
        
        diagnostico_content = None
        diagnostico_filename = None
        
        if diagnostico_file and diagnostico_file.filename:
            if not ProfeGoUtils.validar_extension(diagnostico_file.filename):
                raise HTTPException(
                    status_code=400,
                    detail=f"Tipo de archivo no permitido para diagnÃ³stico: {diagnostico_file.filename}"
                )
            
            diagnostico_content = await diagnostico_file.read()
            if len(diagnostico_content) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail="El archivo de diagnÃ³stico excede el lÃ­mite de 15MB"
                )
            diagnostico_filename = diagnostico_file.filename
        
        logger.info(f"âœ… Archivos validados")
        
        # ========== PROCESAMIENTO OCR ==========
        
        logger.info("ðŸ“„ Extrayendo texto del plan de estudios...")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(plan_file.filename).suffix) as tmp_plan:
            tmp_plan.write(plan_content)
            tmp_plan_path = tmp_plan.name
        
        try:
            plan_result = get_text_only(tmp_plan_path)
            
            if not plan_result['success'] or not plan_result['text']:
                raise HTTPException(
                    status_code=400,
                    detail=f"No se pudo extraer texto del plan"
                )
            
            plan_text = plan_result['text']
            logger.info(f"âœ… Texto extraÃ­do del plan: {len(plan_text)} caracteres")
            
            diagnostico_text = None
            
            if diagnostico_content:
                logger.info("ðŸ“„ Extrayendo texto del diagnÃ³stico...")
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(diagnostico_filename).suffix) as tmp_diag:
                    tmp_diag.write(diagnostico_content)
                    tmp_diag_path = tmp_diag.name
                
                try:
                    diagnostico_result = get_text_only(tmp_diag_path)
                    
                    if diagnostico_result['success'] and diagnostico_result['text']:
                        diagnostico_text = diagnostico_result['text']
                        logger.info(f"âœ… Texto extraÃ­do del diagnÃ³stico: {len(diagnostico_text)} caracteres")
                
                finally:
                    if os.path.exists(tmp_diag_path):
                        os.remove(tmp_diag_path)
            
        finally:
            if os.path.exists(tmp_plan_path):
                os.remove(tmp_plan_path)
        
        # ========== RECUPERACIÃ“N RAG - CON ACTIVIDADES ==========
        
        retrieved_docs = {'cuentos': [], 'canciones': [], 'actividades': []}
        rag_context_text = ""
        
        if rag_system is not None:
            logger.info("ðŸ” Recuperando documentos de la biblioteca RAG...")
            
            try:
                query_text = plan_text
                if diagnostico_text:
                    query_text = f"{plan_text}\n\n{diagnostico_text}"
                
                query_embedding = rag_system.embeddings.embed_query(query_text)
                
                # Buscar cuentos
                logger.info("ðŸ“– Buscando cuentos relevantes...")
                cuentos_results = rag_system.vector_store.query(
                    query_embedding=query_embedding,
                    n_results=5,
                    filter_metadata={'document_type': 'cuento'}
                )
                
                for doc, metadata, distance in zip(
                    cuentos_results['documents'],
                    cuentos_results['metadatas'],
                    cuentos_results['distances']
                ):
                    retrieved_docs['cuentos'].append({
                        'text': doc,
                        'metadata': metadata,
                        'similarity': 1 - distance
                    })
                
                logger.info(f"âœ… {len(retrieved_docs['cuentos'])} cuentos recuperados")
                
                # Buscar canciones
                logger.info("ðŸŽµ Buscando canciones relevantes...")
                canciones_results = rag_system.vector_store.query(
                    query_embedding=query_embedding,
                    n_results=5,
                    filter_metadata={'document_type': 'cancion'}
                )
                
                for doc, metadata, distance in zip(
                    canciones_results['documents'],
                    canciones_results['metadatas'],
                    canciones_results['distances']
                ):
                    retrieved_docs['canciones'].append({
                        'text': doc,
                        'metadata': metadata,
                        'similarity': 1 - distance
                    })
                
                logger.info(f"âœ… {len(retrieved_docs['canciones'])} canciones recuperadas")
                
                # â­ BUSCAR ACTIVIDADES
                logger.info("ðŸŽ¯ Buscando actividades relevantes...")
                actividades_results = rag_system.vector_store.query(
                    query_embedding=query_embedding,
                    n_results=5,
                    filter_metadata={'document_type': 'actividad'}
                )
                
                for doc, metadata, distance in zip(
                    actividades_results['documents'],
                    actividades_results['metadatas'],
                    actividades_results['distances']
                ):
                    retrieved_docs['actividades'].append({
                        'text': doc,
                        'metadata': metadata,
                        'similarity': 1 - distance
                    })
                
                logger.info(f"âœ… {len(retrieved_docs['actividades'])} actividades recuperadas")
                
                # CONSTRUIR CONTEXTO RAG PARA GEMINI
                rag_context_parts = []
                
                if retrieved_docs['cuentos']:
                    rag_context_parts.append("\n\n# ðŸ“– CUENTOS DISPONIBLES EN LA BIBLIOTECA:")
                    for idx, cuento in enumerate(retrieved_docs['cuentos'], 1):
                        filename = cuento['metadata'].get('filename', 'Desconocido')
                        similitud = cuento['similarity'] * 100
                        texto = cuento['text'][:500]
                        
                        rag_context_parts.append(f"""
## Cuento {idx}: {filename}
**Relevancia:** {similitud:.1f}%
**Contenido:**
{texto}
""")
                
                if retrieved_docs['canciones']:
                    rag_context_parts.append("\n\n# ðŸŽµ CANCIONES DISPONIBLES EN LA BIBLIOTECA:")
                    for idx, cancion in enumerate(retrieved_docs['canciones'], 1):
                        filename = cancion['metadata'].get('filename', 'Desconocido')
                        similitud = cancion['similarity'] * 100
                        texto = cancion['text'][:500]
                        
                        rag_context_parts.append(f"""
## CanciÃ³n {idx}: {filename}
**Relevancia:** {similitud:.1f}%
**Contenido:**
{texto}
""")
                
                # â­ AGREGAR ACTIVIDADES AL CONTEXTO
                if retrieved_docs['actividades']:
                    rag_context_parts.append("\n\n# ðŸŽ¯ ACTIVIDADES DIDÃCTICAS DISPONIBLES EN LA BIBLIOTECA:")
                    for idx, actividad in enumerate(retrieved_docs['actividades'], 1):
                        filename = actividad['metadata'].get('filename', 'Desconocido')
                        similitud = actividad['similarity'] * 100
                        texto = actividad['text'][:800]  # MÃ¡s caracteres para actividades
                        
                        rag_context_parts.append(f"""
## Actividad {idx}: {filename}
**Relevancia:** {similitud:.1f}%
**Contenido completo:**
{texto}
""")
                
                rag_context_text = "\n".join(rag_context_parts)
                
                logger.info(f"âœ… Contexto RAG construido: {len(rag_context_text)} caracteres")
                
            except Exception as e:
                logger.warning(f"âš ï¸ Error en RAG, continuando sin Ã©l: {e}")
                rag_context_text = ""
        
        # ========== GENERACIÃ“N CON GEMINI - USANDO CONTEXTO RAG CON ACTIVIDADES ==========
        
        logger.info("ðŸ¤– Generando plan con Gemini AI + contexto RAG (incluye actividades)...")
        
        enriched_plan_text = plan_text
        if rag_context_text:
            enriched_plan_text = f"""
{plan_text}

---

# RECURSOS EDUCATIVOS DISPONIBLES EN LA BIBLIOTECA DIGITAL

{rag_context_text}

---

**INSTRUCCIÃ“N IMPORTANTE PARA LA GENERACIÃ“N:**
Los recursos anteriores (cuentos, canciones y actividades) estÃ¡n VERIFICADOS y DISPONIBLES en la biblioteca.
Al generar el plan:
1. **PRIORIZA** estos recursos en la secciÃ³n "recursos_educativos"
2. Menciona sus tÃ­tulos EXACTOS como aparecen arriba
3. Marca estos recursos como "RECURSO REAL" y "GRATUITO"
4. Indica que estÃ¡n "Disponibles en la biblioteca digital"
5. Integra estos recursos en las actividades cuando sea relevante

**INSTRUCCIONES ESPECIALES PARA ACTIVIDADES:**
- Las actividades de la biblioteca tienen estructura completa con: tÃ­tulo, lÃ­nea de trabajo, Ã¡mbito, organizaciÃ³n, aprendizajes esperados, materiales, desarrollo paso a paso, y sugerencias
- Si una actividad es perfecta para un mÃ³dulo â†’ inclÃºyela completa o adÃ¡ptala en "actividades_desarrollo"
- Marca con "basada_en_actividad_biblioteca": "SI" y especifica el nombre del archivo en "fuente_actividad"
- InclÃºyelas tambiÃ©n en "actividades_complementarias" de la secciÃ³n "recursos_educativos"
"""
            logger.info("âœ… Plan enriquecido con contexto RAG (incluye actividades)")
        
        # Generar con Gemini
        resultado_gemini = await generar_plan_estudio(
            plan_text=enriched_plan_text,
            diagnostico_text=diagnostico_text
        )
        
        if not resultado_gemini['success']:
            raise HTTPException(
                status_code=500,
                detail=f"Error generando plan con IA: {resultado_gemini.get('error')}"
            )
        
        plan_data = resultado_gemini['plan']
        
        logger.info(f"âœ… Plan generado: {plan_data['nombre_plan']}")
        
        # ========== AGREGAR METADATA RAG CON ACTIVIDADES ==========
        
        plan_id = f"plan_{uuid.uuid4().hex[:12]}_{int(datetime.now().timestamp())}"
        
        plan_data['plan_id'] = plan_id
        plan_data['usuario'] = user_email
        plan_data['fecha_generacion'] = datetime.now().isoformat()
        plan_data['archivos_originales'] = {
            'plan': plan_file.filename,
            'diagnostico': diagnostico_filename
        }
        
        # â­ GUARDAR METADATA RAG CON ACTIVIDADES
        plan_data['rag_metadata'] = {
            'recursos_recuperados': {
                'cuentos': [
                    {
                        'nombre': c['metadata'].get('filename', ''),
                        'similitud': round(c['similarity'], 3)
                    }
                    for c in retrieved_docs['cuentos']
                ],
                'canciones': [
                    {
                        'nombre': c['metadata'].get('filename', ''),
                        'similitud': round(c['similarity'], 3)
                    }
                    for c in retrieved_docs['canciones']
                ],
                'actividades': [
                    {
                        'nombre': a['metadata'].get('filename', ''),
                        'similitud': round(a['similarity'], 3)
                    }
                    for a in retrieved_docs['actividades']
                ]
            },
            'total_recuperado': len(retrieved_docs['cuentos']) + len(retrieved_docs['canciones']) + len(retrieved_docs['actividades']),
            'contexto_rag_chars': len(rag_context_text),
            'rag_usado': len(rag_context_text) > 0
        }
        
        logger.info(f"ðŸ“Š Metadata RAG: {plan_data['rag_metadata']['total_recuperado']} recursos (incluye actividades)")
        
        # ========== GUARDAR EN GCS ==========
        
        plan_json = json.dumps(plan_data, indent=2, ensure_ascii=False)
        plan_json_bytes = plan_json.encode('utf-8')
        
        resultado_guardado = gcs_storage.subir_archivo_desde_bytes(
            contenido=plan_json_bytes,
            email=user_email,
            nombre_archivo=f"{plan_id}.json",
            es_procesado=True
        )
        
        if resultado_guardado['success']:
            logger.info(f"âœ… Plan guardado en GCS con metadata RAG (incluye actividades)")
        
        # Subir archivos originales
        gcs_storage.subir_archivo_desde_bytes(
            contenido=plan_content,
            email=user_email,
            nombre_archivo=plan_file.filename,
            es_procesado=False
        )
        
        if diagnostico_content:
            gcs_storage.subir_archivo_desde_bytes(
                contenido=diagnostico_content,
                email=user_email,
                nombre_archivo=diagnostico_filename,
                es_procesado=False
            )
        
        # ========== RETORNAR RESULTADO ==========
        
        processing_time = time.time() - start_time
        logger.info(f"â±ï¸ Tiempo total: {processing_time:.2f}s")
        logger.info(f"ðŸŽ‰ Plan generado exitosamente con RAG (incluye actividades)")
        
        return PlanResponse(
            success=True,
            plan_id=plan_id,
            plan_data=plan_data,
            processing_time=processing_time
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error generando plan: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Error inesperado: {str(e)}"
        )

# ============================================================================
# OTRAS RUTAS DE PLANES
# ============================================================================

@app.get("/api/plans/list")
async def list_plans(
    current_user: dict = Depends(get_current_user)
):
    """Lista todos los planes generados del usuario"""
    user_email = current_user["email"]
    
    try:
        archivos_procesados = gcs_storage.listar_archivos(user_email, "processed")
        planes = []
        
        for archivo in archivos_procesados:
            if archivo['name'].startswith('plan_') and archivo['name'].endswith('.json'):
                contenido = gcs_storage.obtener_archivo_bytes(
                    email=user_email,
                    nombre_archivo=archivo['name'],
                    es_procesado=True
                )
                
                if contenido:
                    try:
                        plan_data = json.loads(contenido.decode('utf-8'))
                        
                        planes.append({
                            'plan_id': plan_data.get('plan_id'),
                            'nombre_plan': plan_data.get('nombre_plan'),
                            'grado': plan_data.get('grado'),
                            'campo_formativo_principal': plan_data.get('campo_formativo_principal'),
                            'ejes_articuladores_generales': plan_data.get('ejes_articuladores_generales', []),
                            'edad_aprox': plan_data.get('edad_aprox'),
                            'duracion_total': plan_data.get('duracion_total'),
                            'materia': plan_data.get('materia'),
                            'num_modulos': plan_data.get('num_modulos', len(plan_data.get('modulos', []))),
                            'fecha_generacion': plan_data.get('fecha_generacion'),
                            'tiene_diagnostico': plan_data.get('tiene_diagnostico', False),
                            'archivos_originales': plan_data.get('archivos_originales', {}),
                            'generado_con': plan_data.get('generado_con'),
                            'modelo': plan_data.get('modelo')
                        })
                    except json.JSONDecodeError:
                        logger.warning(f"âš ï¸ No se pudo parsear el plan: {archivo['name']}")
        
        planes.sort(key=lambda x: x.get('fecha_generacion', ''), reverse=True)
        
        return {
            'success': True,
            'planes': planes,
            'total': len(planes)
        }
        
    except Exception as e:
        logger.error(f"âŒ Error listando planes: {e}")
        raise HTTPException(status_code=500, detail=f"Error listando planes: {str(e)}")

@app.get("/api/plans/{plan_id}")
async def get_plan_detail(
    plan_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Obtiene el detalle completo de un plan especÃ­fico"""
    user_email = current_user["email"]
    
    try:
        filename = f"{plan_id}.json"
        
        contenido = gcs_storage.obtener_archivo_bytes(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=True
        )
        
        if not contenido:
            raise HTTPException(status_code=404, detail="Plan no encontrado")
        
        plan_data = json.loads(contenido.decode('utf-8'))
        
        return {
            'success': True,
            'plan': plan_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error obteniendo plan: {e}")
        raise HTTPException(status_code=500, detail=f"Error obteniendo plan: {str(e)}")

@app.delete("/api/plans/{plan_id}")
@limiter.limit("10/minute")
async def delete_plan(
    request: Request,
    plan_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Elimina un plan generado"""
    user_email = current_user["email"]
    
    try:
        filename = f"{plan_id}.json"
        
        resultado = gcs_storage.eliminar_archivo(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=True
        )
        
        if not resultado['success']:
            raise HTTPException(status_code=404, detail="Plan no encontrado")
        
        return {
            'success': True,
            'message': 'Plan eliminado correctamente'
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error eliminando plan: {e}")
        raise HTTPException(status_code=500, detail=f"Error eliminando plan: {str(e)}")
    
# ============================================================================
# GENERACIÃ“N DE DOCUMENTOS WORD
# ============================================================================

def generar_documento_word(plan_data: Dict) -> io.BytesIO:
    """
    Genera un documento Word profesional a partir de los datos del plan
    VERSIÃ“N CON SOPORTE PARA ACTIVIDADES COMPLEMENTARIAS
    """
    doc = Document()
    
    # Configurar estilos del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    p_format = style.paragraph_format
    p_format.line_spacing = 1.5
    
    # ========== PORTADA ==========
    portada = doc.add_heading(plan_data.get('nombre_plan', 'Plan Educativo'), 0)
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    portada.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    # InformaciÃ³n general
    info_general = doc.add_paragraph()
    info_general.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if plan_data.get('grado'):
        run = info_general.add_run(f"Grado: {plan_data['grado']}\n")
        run.bold = True
        run.font.size = Pt(12)
    
    if plan_data.get('edad_aprox'):
        run = info_general.add_run(f"Edad aproximada: {plan_data['edad_aprox']}\n")
        run.font.size = Pt(12)
    
    if plan_data.get('duracion_total'):
        run = info_general.add_run(f"DuraciÃ³n total: {plan_data['duracion_total']}\n")
        run.font.size = Pt(12)

    if plan_data.get('fecha_generacion'):
        fecha = datetime.fromisoformat(plan_data['fecha_generacion']).strftime('%d/%m/%Y %H:%M')
        run = info_general.add_run(f"Generado: {fecha}\n")
        run.font.size = Pt(12)
        run.italic = True
    
    doc.add_paragraph()
    
    # Campo formativo y ejes
    if plan_data.get('campo_formativo_principal'):
        p = doc.add_paragraph()
        p.add_run('Campo Formativo Principal: ').bold = True
        p.add_run(plan_data['campo_formativo_principal'])
    
    if plan_data.get('ejes_articuladores_generales'):
        p = doc.add_paragraph()
        p.add_run('Ejes Articuladores: ').bold = True
        p.add_run(', '.join(plan_data['ejes_articuladores_generales']))
    
    doc.add_page_break()
    
    # ========== MÃ“DULOS ==========
    doc.add_heading('â—ˆ MÃ³dulos del Plan', 1)
    
    modulos = plan_data.get('modulos', [])
    
    for idx, modulo in enumerate(modulos, 1):
        # Encabezado del mÃ³dulo
        modulo_heading = doc.add_heading(f"MÃ³dulo {modulo.get('numero', idx)}: {modulo.get('nombre', '')}", 2)
        modulo_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
        # InformaciÃ³n del mÃ³dulo
        if modulo.get('campo_formativo'):
            p = doc.add_paragraph()
            p.add_run('â—‡ Campo Formativo: ').bold = True
            p.add_run(modulo['campo_formativo'])
        
        if modulo.get('ejes_articuladores'):
            p = doc.add_paragraph()
            p.add_run('â—‡ Ejes Articuladores: ').bold = True
            p.add_run(', '.join(modulo['ejes_articuladores']))
        
        if modulo.get('aprendizaje_esperado'):
            p = doc.add_paragraph()
            p.add_run('â—‡ Aprendizaje Esperado: ').bold = True
            p.add_run(modulo['aprendizaje_esperado'])
        
        if modulo.get('tiempo_estimado'):
            p = doc.add_paragraph()
            p.add_run('â—‡ Tiempo Estimado: ').bold = True
            p.add_run(modulo['tiempo_estimado'])
        
        doc.add_paragraph()
        
        # Actividad de inicio
        if modulo.get('actividad_inicio'):
            doc.add_heading('â—† Actividad de Inicio', 3)
            inicio = modulo['actividad_inicio']
            
            p = doc.add_paragraph()
            p.add_run('Nombre: ').bold = True
            p.add_run(inicio.get('nombre', ''))
            
            p = doc.add_paragraph()
            p.add_run('DescripciÃ³n: ').bold = True
            p.add_run(inicio.get('descripcion', ''))
            
            if inicio.get('duracion'):
                p = doc.add_paragraph()
                p.add_run('DuraciÃ³n: ').bold = True
                p.add_run(inicio['duracion'])
            
            if inicio.get('materiales'):
                p = doc.add_paragraph()
                p.add_run('Materiales: ').bold = True
                materiales = inicio['materiales'] if isinstance(inicio['materiales'], list) else [inicio['materiales']]
                p.add_run(', '.join(materiales))
            
            if inicio.get('organizacion'):
                p = doc.add_paragraph()
                p.add_run('OrganizaciÃ³n: ').bold = True
                p.add_run(inicio['organizacion'])
        
        # Actividades de desarrollo
        if modulo.get('actividades_desarrollo'):
            doc.add_heading('â—† Actividades de Desarrollo', 3)
            
            for act_idx, actividad in enumerate(modulo['actividades_desarrollo'], 1):
                doc.add_heading(f"Actividad {act_idx}: {actividad.get('nombre', '')}", 4)
                
                if actividad.get('tipo'):
                    p = doc.add_paragraph()
                    p.add_run('Tipo: ').bold = True
                    p.add_run(actividad['tipo'])
                
                # â­ NUEVO: Mostrar si estÃ¡ basada en la biblioteca
                if actividad.get('basada_en_actividad_biblioteca') == 'SI':
                    p = doc.add_paragraph()
                    p.add_run('ðŸ“š Basada en biblioteca: ').bold = True
                    p.add_run('SÃ')
                    if actividad.get('fuente_actividad'):
                        p = doc.add_paragraph()
                        p.add_run('ðŸ“„ Fuente: ').bold = True
                        p.add_run(actividad['fuente_actividad'])
                
                if actividad.get('descripcion'):
                    p = doc.add_paragraph()
                    p.add_run('DescripciÃ³n: ').bold = True
                    p.add_run(actividad['descripcion'])
                
                if actividad.get('duracion'):
                    p = doc.add_paragraph()
                    p.add_run('DuraciÃ³n: ').bold = True
                    p.add_run(actividad['duracion'])
                
                if actividad.get('organizacion'):
                    p = doc.add_paragraph()
                    p.add_run('OrganizaciÃ³n: ').bold = True
                    p.add_run(actividad['organizacion'])
                
                if actividad.get('materiales'):
                    p = doc.add_paragraph()
                    p.add_run('Materiales: ').bold = True
                    materiales = actividad['materiales'] if isinstance(actividad['materiales'], list) else [actividad['materiales']]
                    p.add_run(', '.join(materiales))
                
                if actividad.get('aspectos_a_observar'):
                    p = doc.add_paragraph()
                    p.add_run('Aspectos a observar: ').bold = True
                    p.add_run(actividad['aspectos_a_observar'])
                
                doc.add_paragraph()
        
        # Actividad de cierre
        if modulo.get('actividad_cierre'):
            doc.add_heading('â—† Actividad de Cierre', 3)
            cierre = modulo['actividad_cierre']
            
            p = doc.add_paragraph()
            p.add_run('Nombre: ').bold = True
            p.add_run(cierre.get('nombre', ''))
            
            p = doc.add_paragraph()
            p.add_run('DescripciÃ³n: ').bold = True
            p.add_run(cierre.get('descripcion', ''))
            
            if cierre.get('duracion'):
                p = doc.add_paragraph()
                p.add_run('DuraciÃ³n: ').bold = True
                p.add_run(cierre['duracion'])
            
            if cierre.get('preguntas_guia'):
                p = doc.add_paragraph()
                p.add_run('Preguntas guÃ­a:').bold = True
                for pregunta in cierre['preguntas_guia']:
                    doc.add_paragraph(f" {pregunta}", style='List Bullet')
        
        # InformaciÃ³n adicional del mÃ³dulo
        if modulo.get('consejos_maestra'):
            doc.add_heading('â—‡ Consejos para el Docente', 3)
            doc.add_paragraph(modulo['consejos_maestra'])
        
        if modulo.get('variaciones'):
            doc.add_heading('â—‡ Variaciones', 3)
            doc.add_paragraph(modulo['variaciones'])
        
        if modulo.get('vinculo_familia'):
            doc.add_heading('â—‡ VÃ­nculo con la Familia', 3)
            doc.add_paragraph(modulo['vinculo_familia'])
        
        if modulo.get('evaluacion'):
            doc.add_heading('â—‡ EvaluaciÃ³n', 3)
            doc.add_paragraph(modulo['evaluacion'])
        
        # Separador entre mÃ³dulos
        if idx < len(modulos):
            doc.add_page_break()
    
    # ========== RECURSOS EDUCATIVOS ==========
    if plan_data.get('recursos_educativos'):
        doc.add_page_break()
        doc.add_heading('â—ˆ Recursos Educativos', 1)
        recursos = plan_data['recursos_educativos']
        
        if recursos.get('materiales_generales'):
            doc.add_heading('â—‡ Materiales Generales', 2)
            for material in recursos['materiales_generales']:
                doc.add_paragraph(f" {material}", style='List Bullet')
        
        if recursos.get('cuentos_recomendados'):
            doc.add_heading('â—‡ Cuentos Recomendados', 2)
            for cuento in recursos['cuentos_recomendados']:
                p = doc.add_paragraph()
                p.add_run(f"â€¢ {cuento.get('titulo', '')}: ").bold = True
                
                detalles = []
                if cuento.get('autor'):
                    detalles.append(f"Autor: {cuento['autor']}")
                if cuento.get('tipo'):
                    detalles.append(f"Tipo: {cuento['tipo']}")
                if cuento.get('acceso'):
                    detalles.append(f"Acceso: {cuento['acceso']}")
                if cuento.get('disponibilidad'):
                    detalles.append(f"Disponibilidad: {cuento['disponibilidad']}")
                
                p.add_run(' | '.join(detalles))
                
                if cuento.get('descripcion_breve'):
                    doc.add_paragraph(f"  {cuento['descripcion_breve']}", style='List Bullet 2')
        
        if recursos.get('canciones_recomendadas'):
            doc.add_heading('â—‡ Canciones Recomendadas', 2)
            for cancion in recursos['canciones_recomendadas']:
                p = doc.add_paragraph()
                p.add_run(f"â€¢ {cancion.get('titulo', '')}: ").bold = True
                
                detalles = []
                if cancion.get('tipo'):
                    detalles.append(f"Tipo: {cancion['tipo']}")
                if cancion.get('acceso'):
                    detalles.append(f"Acceso: {cancion['acceso']}")
                if cancion.get('disponibilidad'):
                    detalles.append(f"Disponibilidad: {cancion['disponibilidad']}")
                
                p.add_run(' | '.join(detalles))
                
                if cancion.get('uso_sugerido'):
                    doc.add_paragraph(f"  Uso sugerido: {cancion['uso_sugerido']}", style='List Bullet 2')
        
        # â­ NUEVA SECCIÃ“N: ACTIVIDADES COMPLEMENTARIAS
        if recursos.get('actividades_complementarias'):
            doc.add_heading('ðŸŽ¯ Actividades Complementarias', 2)
            for actividad in recursos['actividades_complementarias']:
                p = doc.add_paragraph()
                p.add_run(f"â€¢ {actividad.get('titulo', '')}: ").bold = True
                
                detalles = []
                if actividad.get('linea_trabajo'):
                    detalles.append(f"LÃ­nea: {actividad['linea_trabajo']}")
                if actividad.get('ambito'):
                    detalles.append(f"Ãmbito: {actividad['ambito']}")
                if actividad.get('organizacion'):
                    detalles.append(f"OrganizaciÃ³n: {actividad['organizacion']}")
                if actividad.get('tipo'):
                    detalles.append(f"Tipo: {actividad['tipo']}")
                if actividad.get('acceso'):
                    detalles.append(f"Acceso: {actividad['acceso']}")
                
                p.add_run(' | '.join(detalles))
                
                if actividad.get('descripcion_breve'):
                    doc.add_paragraph(f"  {actividad['descripcion_breve']}", style='List Bullet 2')
                
                if actividad.get('materiales_necesarios'):
                    p_materiales = doc.add_paragraph(style='List Bullet 2')
                    p_materiales.add_run('  Materiales: ').bold = True
                    p_materiales.add_run(', '.join(actividad['materiales_necesarios']))
    
    # ========== RECOMENDACIONES DE AMBIENTE ==========
    if plan_data.get('recomendaciones_ambiente'):
        doc.add_heading('â—‡ Recomendaciones para el Ambiente', 2)
        doc.add_paragraph(plan_data['recomendaciones_ambiente'])
    
    # ========== VINCULACIÃ“N CURRICULAR ==========
    if plan_data.get('vinculacion_curricular'):
        doc.add_heading('â—‡ VinculaciÃ³n Curricular', 2)
        vinculacion = plan_data['vinculacion_curricular']
        
        if vinculacion.get('campo_formativo_principal'):
            p = doc.add_paragraph()
            p.add_run('Campo Formativo Principal: ').bold = True
            p.add_run(vinculacion['campo_formativo_principal'])
        
        if vinculacion.get('campos_secundarios'):
            p = doc.add_paragraph()
            p.add_run('Campos Secundarios: ').bold = True
            p.add_run(', '.join(vinculacion['campos_secundarios']))
        
        if vinculacion.get('ejes_transversales'):
            p = doc.add_paragraph()
            p.add_run('Ejes Transversales: ').bold = True
            p.add_run(', '.join(vinculacion['ejes_transversales']))
        
        if vinculacion.get('aprendizajes_clave'):
            doc.add_heading('Aprendizajes Clave:', 3)
            for aprendizaje in vinculacion['aprendizajes_clave']:
                doc.add_paragraph(f" {aprendizaje}", style='List Bullet')
    
    # Guardar en memoria
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

@app.get("/api/plans/{plan_id}/download")
async def download_plan_word(
    plan_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Descarga el plan como documento Word (.docx) profesional"""
    user_email = current_user["email"]
    
    try:
        filename = f"{plan_id}.json"
        
        contenido = gcs_storage.obtener_archivo_bytes(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=True
        )
        
        if not contenido:
            raise HTTPException(status_code=404, detail="Plan no encontrado")
        
        plan_data = json.loads(contenido.decode('utf-8'))
        
        # Generar documento Word
        docx_bytes = generar_documento_word(plan_data)
        
        # Crear nombre de archivo seguro
        nombre_plan = plan_data.get('nombre_plan', 'Plan_Educativo')
        nombre_archivo = f"{nombre_plan.replace(' ', '_').replace('/', '_')}.docx"
        
        # Retornar archivo Word
        return StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={nombre_archivo}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error generando documento Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando documento: {str(e)}")
    

# ============================================================================
# CLASE RAGAnalyzer CON SOPORTE PARA ACTIVIDADES
# ============================================================================

class RAGAnalyzer:
    """
    Analizador avanzado de similitud semÃ¡ntica entre planes y recursos RAG
    VERSIÃ“N CON SOPORTE PARA ACTIVIDADES
    """
    
    def __init__(self, rag_system):
        self.rag_system = rag_system
    
    def analyze_plan_rag_match(
        self,
        plan_data: Dict,
        retrieved_docs: Dict,
        threshold: float = 0.65
    ) -> Dict:
        """
        Analiza la similitud entre el plan generado y los recursos RAG
        INCLUYE ACTIVIDADES
        """
        analisis = {
            'similitud_general': 0.0,
            'recursos_altamente_relevantes': [],
            'recursos_por_modulo': [],
            'metricas_rag': {
                'total_recursos_rag': 0,
                'recursos_utilizados': 0,
                'porcentaje_uso_rag': 0.0,
                'similitud_promedio': 0.0,
                'actividades_biblioteca_usadas': 0  # â­ NUEVO
            },
            'recomendaciones_adicionales': []
        }
        
        # Extraer texto completo del plan
        plan_text = self._extract_plan_text(plan_data)
        plan_text_lower = plan_text.lower()
        
        # Combinar todos los recursos
        cuentos = retrieved_docs.get('cuentos', [])
        canciones = retrieved_docs.get('canciones', [])
        actividades = retrieved_docs.get('actividades', [])  # â­ NUEVO
        
        analisis['metricas_rag']['total_recursos_rag'] = len(cuentos) + len(canciones) + len(actividades)
        
        # AnÃ¡lisis mejorado: Buscar coincidencias semÃ¡nticas
        recursos_encontrados = []
        
        # Procesar cuentos
        for cuento in cuentos:
            filename = cuento['metadata'].get('filename', '')
            clean_name = Path(filename).stem.replace('_', ' ').lower()
            similitud = cuento.get('similarity', 0)
            
            usado = False
            keywords = clean_name.split()[:3]
            
            for keyword in keywords:
                if len(keyword) > 3 and keyword in plan_text_lower:
                    usado = True
                    break
            
            # TambiÃ©n revisar si estÃ¡ en la secciÃ³n de recursos
            recursos_plan = plan_data.get('recursos_educativos', {})
            cuentos_recomendados = recursos_plan.get('cuentos_recomendados', [])
            
            for cuento_rec in cuentos_recomendados:
                titulo = cuento_rec.get('titulo', '').lower()
                if any(kw in titulo for kw in keywords if len(kw) > 3):
                    usado = True
                    break
            
            if usado or similitud >= 0.65:
                recurso_info = self._format_recurso(cuento, 'cuento', similitud)
                recursos_encontrados.append(recurso_info)
                if usado:
                    analisis['metricas_rag']['recursos_utilizados'] += 1
        
        # Procesar canciones
        for cancion in canciones:
            filename = cancion['metadata'].get('filename', '')
            clean_name = Path(filename).stem.replace('_', ' ').lower()
            similitud = cancion.get('similarity', 0)
            
            usado = False
            keywords = clean_name.split()[:4]
            
            for keyword in keywords:
                if len(keyword) > 4 and keyword in plan_text_lower:
                    usado = True
                    break
            
            recursos_plan = plan_data.get('recursos_educativos', {})
            canciones_recomendadas = recursos_plan.get('canciones_recomendadas', [])
            
            for cancion_rec in canciones_recomendadas:
                titulo = cancion_rec.get('titulo', '').lower()
                if any(kw in titulo for kw in keywords if len(kw) > 3):
                    usado = True
                    break
            
            if usado or similitud >= 0.65:
                recurso_info = self._format_recurso(cancion, 'cancion', similitud)
                recursos_encontrados.append(recurso_info)
                if usado:
                    analisis['metricas_rag']['recursos_utilizados'] += 1
        
        # â­ PROCESAR ACTIVIDADES
        for actividad in actividades:
            filename = actividad['metadata'].get('filename', '')
            clean_name = Path(filename).stem.replace('_', ' ').lower()
            similitud = actividad.get('similarity', 0)
            
            usado = False
            
            # Buscar en actividades_desarrollo
            for modulo in plan_data.get('modulos', []):
                for act in modulo.get('actividades_desarrollo', []):
                    # Si estÃ¡ marcada como basada en biblioteca
                    if act.get('basada_en_actividad_biblioteca') == 'SI':
                        fuente = act.get('fuente_actividad', '').lower()
                        if filename.lower() in fuente or clean_name in act.get('nombre', '').lower():
                            usado = True
                            analisis['metricas_rag']['actividades_biblioteca_usadas'] += 1
                            break
                if usado:
                    break
            
            # Buscar en actividades_complementarias
            if not usado:
                recursos_plan = plan_data.get('recursos_educativos', {})
                actividades_complementarias = recursos_plan.get('actividades_complementarias', [])
                
                for act_comp in actividades_complementarias:
                    titulo = act_comp.get('titulo', '').lower()
                    keywords = clean_name.split()[:3]
                    if any(kw in titulo for kw in keywords if len(kw) > 3):
                        usado = True
                        analisis['metricas_rag']['actividades_biblioteca_usadas'] += 1
                        break
            
            if usado or similitud >= 0.65:
                recurso_info = self._format_recurso(actividad, 'actividad', similitud)
                recursos_encontrados.append(recurso_info)
                if usado:
                    analisis['metricas_rag']['recursos_utilizados'] += 1
        
        # Ordenar por similitud
        recursos_encontrados.sort(key=lambda x: x['similitud_porcentaje'], reverse=True)
        analisis['recursos_altamente_relevantes'] = recursos_encontrados
        
        # Calcular porcentaje basado en recursos_educativos
        recursos_plan = plan_data.get('recursos_educativos', {})
        total_cuentos_plan = len(recursos_plan.get('cuentos_recomendados', []))
        total_canciones_plan = len(recursos_plan.get('canciones_recomendadas', []))
        total_actividades_plan = len(recursos_plan.get('actividades_complementarias', []))  # â­ NUEVO
        total_recursos_plan = total_cuentos_plan + total_canciones_plan + total_actividades_plan
        
        if total_recursos_plan > 0:
            # Contar cuÃ¡ntos recursos del plan tienen origen RAG
            recursos_rag_en_plan = 0
            
            for cuento_plan in recursos_plan.get('cuentos_recomendados', []):
                if cuento_plan.get('tipo') == 'RECURSO REAL':
                    recursos_rag_en_plan += 1
            
            for cancion_plan in recursos_plan.get('canciones_recomendadas', []):
                if cancion_plan.get('tipo') == 'RECURSO REAL':
                    recursos_rag_en_plan += 1
            
            # â­ Contar actividades de la biblioteca
            for actividad_plan in recursos_plan.get('actividades_complementarias', []):
                if actividad_plan.get('tipo') == 'RECURSO REAL':
                    recursos_rag_en_plan += 1
            
            analisis['metricas_rag']['recursos_utilizados'] = max(
                analisis['metricas_rag']['recursos_utilizados'],
                recursos_rag_en_plan
            )
            
            analisis['metricas_rag']['porcentaje_uso_rag'] = round(
                (recursos_rag_en_plan / total_recursos_plan) * 100,
                1
            )
        elif analisis['metricas_rag']['total_recursos_rag'] > 0:
            # Fallback
            analisis['metricas_rag']['porcentaje_uso_rag'] = round(
                (analisis['metricas_rag']['recursos_utilizados'] / 
                 analisis['metricas_rag']['total_recursos_rag']) * 100,
                1
            )
        
        # Calcular similitud promedio
        if recursos_encontrados:
            analisis['metricas_rag']['similitud_promedio'] = round(
                sum(r['similitud_porcentaje'] for r in recursos_encontrados) / len(recursos_encontrados),
                1
            )
        
        # Analizar por mÃ³dulo
        analisis['recursos_por_modulo'] = self._analyze_modules(
            plan_data,
            retrieved_docs,
            threshold
        )
        
        logger.info(f"ðŸ“Š AnÃ¡lisis RAG completado:")
        logger.info(f"   Total recuperado: {analisis['metricas_rag']['total_recursos_rag']}")
        logger.info(f"   Recursos utilizados: {analisis['metricas_rag']['recursos_utilizados']}")
        logger.info(f"   Actividades biblioteca: {analisis['metricas_rag']['actividades_biblioteca_usadas']}")
        logger.info(f"   Porcentaje uso: {analisis['metricas_rag']['porcentaje_uso_rag']}%")
        logger.info(f"   Similitud promedio: {analisis['metricas_rag']['similitud_promedio']}%")
        
        return analisis
    
    def _extract_plan_text(self, plan_data: Dict) -> str:
        """Extrae el texto completo del plan para anÃ¡lisis"""
        texts = []
        
        texts.append(plan_data.get('nombre_plan', ''))
        texts.append(plan_data.get('campo_formativo_principal', ''))
        
        for modulo in plan_data.get('modulos', []):
            texts.append(modulo.get('nombre', ''))
            texts.append(modulo.get('aprendizaje_esperado', ''))
            
            if modulo.get('actividad_inicio'):
                texts.append(modulo['actividad_inicio'].get('descripcion', ''))
            
            for act in modulo.get('actividades_desarrollo', []):
                texts.append(act.get('descripcion', ''))
                # â­ Incluir nombre de actividad y fuente
                texts.append(act.get('nombre', ''))
                texts.append(act.get('fuente_actividad', ''))
            
            if modulo.get('actividad_cierre'):
                texts.append(modulo['actividad_cierre'].get('descripcion', ''))
        
        # Incluir recursos educativos
        recursos = plan_data.get('recursos_educativos', {})
        
        for cuento in recursos.get('cuentos_recomendados', []):
            texts.append(cuento.get('titulo', ''))
            texts.append(cuento.get('descripcion_breve', ''))
        
        for cancion in recursos.get('canciones_recomendadas', []):
            texts.append(cancion.get('titulo', ''))
        
        # â­ Incluir actividades complementarias
        for actividad in recursos.get('actividades_complementarias', []):
            texts.append(actividad.get('titulo', ''))
            texts.append(actividad.get('descripcion_breve', ''))
        
        return ' '.join(filter(None, texts))
    
    def _format_recurso(self, doc: Dict, tipo: str, similitud: float) -> Dict:
        """Formatea un recurso RAG para presentaciÃ³n"""
        metadata = doc.get('metadata', {})
        texto = doc.get('text', '')
        
        # Extraer tÃ­tulo del nombre del archivo
        filename = metadata.get('filename', 'Recurso desconocido')
        titulo = Path(filename).stem.replace('_', ' ').title()
        
        return {
            'titulo': titulo,
            'tipo': tipo,
            'fuente': 'RECURSO REAL',
            'similitud_porcentaje': round(similitud * 100, 1),
            'similitud_nivel': self._get_similarity_level(similitud),
            'contenido_completo': texto,
            'fragmento': texto[:300] + '...' if len(texto) > 300 else texto,
            'filename': filename,
            'acceso': 'GRATUITO',
            'markdown_formato': f"""---
### ðŸ“š {titulo}

**Tipo:** {tipo.upper()}  
**Disponibilidad:** Biblioteca Digital ProfeGo  
**Acceso:** GRATUITO  
**Similitud con el plan:** {round(similitud * 100, 1)}%

**Contenido:**

{texto[:500]}...

---
"""
        }
    
    def _get_similarity_level(self, similitud: float) -> str:
        """Clasifica el nivel de similitud"""
        if similitud >= 0.75:
            return 'MUY ALTA'
        elif similitud >= 0.60:
            return 'ALTA'
        elif similitud >= 0.48:
            return 'MEDIA'
        elif similitud >= 0.30:
            return 'MEDIA-BAJA'
        else:
            return 'BAJA'
    
    def _analyze_modules(
        self,
        plan_data: Dict,
        retrieved_docs: Dict,
        threshold: float
    ) -> List[Dict]:
        """Analiza similitud por mÃ³dulo - INCLUYE ACTIVIDADES"""
        modulos_analisis = []
        
        for modulo in plan_data.get('modulos', []):
            modulo_text = f"{modulo.get('nombre', '')} {modulo.get('aprendizaje_esperado', '')}"
            modulo_text_lower = modulo_text.lower()
            
            if not modulo_text.strip():
                continue
            
            recursos_modulo = []
            
            # Buscar cuentos relacionados
            for cuento in retrieved_docs.get('cuentos', []):
                filename = cuento['metadata'].get('filename', '')
                clean_name = Path(filename).stem.replace('_', ' ')
                similitud = cuento.get('similarity', 0)
                
                keywords = clean_name.lower().split()[:3]
                if any(kw in modulo_text_lower for kw in keywords if len(kw) > 3):
                    recursos_modulo.append({
                        'titulo': clean_name.title(),
                        'tipo': 'cuento',
                        'similitud': round(similitud * 100, 1)
                    })
            
            # Buscar canciones relacionadas
            for cancion in retrieved_docs.get('canciones', []):
                filename = cancion['metadata'].get('filename', '')
                clean_name = Path(filename).stem.replace('_', ' ')
                similitud = cancion.get('similarity', 0)
                
                keywords = clean_name.lower().split()[:3]
                if any(kw in modulo_text_lower for kw in keywords if len(kw) > 3):
                    recursos_modulo.append({
                        'titulo': clean_name.title(),
                        'tipo': 'cancion',
                        'similitud': round(similitud * 100, 1)
                    })
            
            # â­ Buscar actividades relacionadas
            for actividad in retrieved_docs.get('actividades', []):
                filename = actividad['metadata'].get('filename', '')
                clean_name = Path(filename).stem.replace('_', ' ')
                similitud = actividad.get('similarity', 0)
                
                keywords = clean_name.lower().split()[:3]
                if any(kw in modulo_text_lower for kw in keywords if len(kw) > 3):
                    recursos_modulo.append({
                        'titulo': clean_name.title(),
                        'tipo': 'actividad',
                        'similitud': round(similitud * 100, 1)
                    })
            
            if recursos_modulo:
                modulos_analisis.append({
                    'numero': modulo.get('numero', 0),
                    'nombre': modulo.get('nombre', ''),
                    'recursos_relacionados': recursos_modulo[:5]  # Top 5
                })
        
        return modulos_analisis


# Instancia global del analizador
rag_analyzer = None


@app.on_event("startup")
async def initialize_rag_analyzer():
    """Inicializa el analizador RAG al arranque"""
    global rag_analyzer
    
    if rag_system is not None:
        rag_analyzer = RAGAnalyzer(rag_system)
        logger.info("âœ… RAG Analyzer inicializado")


# ============================================================================
# RUTAS DE ANÃLISIS RAG
# ============================================================================

@app.get("/api/plans/{plan_id}/rag-analysis")
async def analyze_plan_rag(
    plan_id: str,
    current_user: dict = Depends(get_current_user)
):
    """
    Endpoint: AnÃ¡lisis completo de similitud RAG (INCLUYE ACTIVIDADES)
    """
    user_email = current_user["email"]
    
    try:
        logger.info(f"ðŸ” Iniciando anÃ¡lisis RAG para plan: {plan_id}")
        
        if rag_analyzer is None or rag_system is None:
            logger.warning("Sistema RAG no disponible")
            return {
                'success': False,
                'message': 'El sistema de anÃ¡lisis RAG no estÃ¡ disponible en este momento.'
            }
        
        # Obtener el plan
        filename = f"{plan_id}.json"
        contenido = gcs_storage.obtener_archivo_bytes(
            email=user_email,
            nombre_archivo=filename,
            es_procesado=True
        )
        
        if not contenido:
            logger.warning(f"Plan no encontrado: {plan_id}")
            raise HTTPException(status_code=404, detail="Plan no encontrado")
        
        plan_data = json.loads(contenido.decode('utf-8'))
        logger.info(f"âœ… Plan cargado: {plan_data.get('nombre_plan')}")
        
        # Verificar metadata RAG
        rag_metadata = plan_data.get('rag_metadata', {})
        
        if not rag_metadata or not rag_metadata.get('recursos_recuperados'):
            logger.warning(f"Plan sin metadata RAG: {plan_id}")
            return {
                'success': False,
                'message': 'Este plan no tiene metadata RAG.',
                'plan_name': plan_data.get('nombre_plan'),
                'sugerencia': 'Genera un nuevo plan para que incluya anÃ¡lisis RAG automÃ¡ticamente.'
            }
        
        # Reconstruir retrieved_docs desde metadata
        retrieved_docs = {
            'cuentos': [],
            'canciones': [],
            'actividades': []  # â­ NUEVO
        }
        
        logger.info("ðŸ“š Reconstruyendo documentos RAG desde metadata...")
        
        recursos_metadata = rag_metadata.get('recursos_recuperados', {})
        
        # Buscar cuentos
        cuentos_dir = Path('./rag_data/cuentos')
        for cuento_meta in recursos_metadata.get('cuentos', []):
            filename_cuento = cuento_meta.get('nombre', '')
            if not filename_cuento:
                continue
                
            cuento_path = cuentos_dir / filename_cuento
            
            if cuento_path.exists():
                try:
                    with open(cuento_path, 'r', encoding='utf-8') as f:
                        contenido_cuento = f.read()
                    
                    retrieved_docs['cuentos'].append({
                        'text': contenido_cuento,
                        'metadata': {
                            'filename': filename_cuento,
                            'document_type': 'cuento'
                        },
                        'similarity': cuento_meta.get('similitud', 0.75)
                    })
                    logger.info(f"âœ… Cuento cargado: {filename_cuento}")
                except Exception as e:
                    logger.warning(f"âš ï¸ Error leyendo cuento {filename_cuento}: {e}")
        
        # Buscar canciones
        canciones_dir = Path('./rag_data/canciones')
        for cancion_meta in recursos_metadata.get('canciones', []):
            filename_cancion = cancion_meta.get('nombre', '')
            if not filename_cancion:
                continue
                
            cancion_path = canciones_dir / filename_cancion
            
            if cancion_path.exists():
                try:
                    with open(cancion_path, 'r', encoding='utf-8') as f:
                        contenido_cancion = f.read()
                    
                    retrieved_docs['canciones'].append({
                        'text': contenido_cancion,
                        'metadata': {
                            'filename': filename_cancion,
                            'document_type': 'cancion'
                        },
                        'similarity': cancion_meta.get('similitud', 0.75)
                    })
                    logger.info(f"âœ… CanciÃ³n cargada: {filename_cancion}")
                except Exception as e:
                    logger.warning(f"âš ï¸ Error leyendo canciÃ³n {filename_cancion}: {e}")
        
        # â­ Buscar actividades
        actividades_dir = Path('./rag_data/actividades')
        for actividad_meta in recursos_metadata.get('actividades', []):
            filename_actividad = actividad_meta.get('nombre', '')
            if not filename_actividad:
                continue
                
            actividad_path = actividades_dir / filename_actividad
            
            if actividad_path.exists():
                try:
                    with open(actividad_path, 'r', encoding='utf-8') as f:
                        contenido_actividad = f.read()
                    
                    retrieved_docs['actividades'].append({
                        'text': contenido_actividad,
                        'metadata': {
                            'filename': filename_actividad,
                            'document_type': 'actividad'
                        },
                        'similarity': actividad_meta.get('similitud', 0.75)
                    })
                    logger.info(f"âœ… Actividad cargada: {filename_actividad}")
                except Exception as e:
                    logger.warning(f"âš ï¸ Error leyendo actividad {filename_actividad}: {e}")
        
        total_recursos = len(retrieved_docs['cuentos']) + len(retrieved_docs['canciones']) + len(retrieved_docs['actividades'])
        logger.info(f"ðŸ“Š Total recursos cargados: {total_recursos} (cuentos: {len(retrieved_docs['cuentos'])}, canciones: {len(retrieved_docs['canciones'])}, actividades: {len(retrieved_docs['actividades'])})")
        
        if total_recursos == 0:
            return {
                'success': False,
                'message': 'No se pudieron cargar los recursos RAG desde el filesystem.',
                'plan_name': plan_data.get('nombre_plan')
            }
        
        # Realizar anÃ¡lisis
        logger.info("ðŸ”¬ Analizando similitud semÃ¡ntica...")
        analisis = rag_analyzer.analyze_plan_rag_match(
            plan_data,
            retrieved_docs,
            threshold=0.50
        )
        
        logger.info(f"âœ… AnÃ¡lisis completado: {analisis['metricas_rag']['porcentaje_uso_rag']}% uso RAG")
        logger.info(f"   Actividades biblioteca usadas: {analisis['metricas_rag']['actividades_biblioteca_usadas']}")
        
        # Generar recursos formateados
        recursos_completos = []
        for recurso in analisis['recursos_altamente_relevantes']:
            recursos_completos.append({
                **recurso,
                'markdown_formato': f"""---
### ðŸ“š Recurso Educativo Relacionado

**TÃ­tulo:** {recurso['titulo']}  
**Tipo:** {recurso['tipo'].upper()}  
**Fuente:** {recurso['fuente']}  
**Similitud con el plan:** {recurso['similitud_porcentaje']}% ({recurso['similitud_nivel']})  
**Acceso:** {recurso['acceso']}

**Contenido completo:**

{recurso['contenido_completo']}

---
"""
            })
        
        return {
            'success': True,
            'plan_id': plan_id,
            'plan_name': plan_data.get('nombre_plan'),
            'analisis': analisis,
            'recursos_completos': recursos_completos
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"âŒ Error analizando RAG: {e}", exc_info=True)
        return {
            'success': False,
            'message': f'Error inesperado al analizar RAG: {str(e)}',
            'error_type': type(e).__name__
        }

# ============================================================================
# RUTAS AUXILIARES Y DE DEBUG
# ============================================================================

@app.get("/api/rag/debug/status")
async def rag_debug_status():
    """Endpoint de debug para verificar estado del sistema RAG"""
    cuentos_path = Path('./rag_data/cuentos')
    canciones_path = Path('./rag_data/canciones')
    actividades_path = Path('./rag_data/actividades')
    
    cuentos_files = list(cuentos_path.glob('**/*.txt')) if cuentos_path.exists() else []
    canciones_files = list(canciones_path.glob('**/*.txt')) if canciones_path.exists() else []
    actividades_files = list(actividades_path.glob('**/*.txt')) if actividades_path.exists() else []
    
    status = {
        'rag_system_initialized': rag_system is not None,
        'rag_analyzer_initialized': rag_analyzer is not None,
        'filesystem': {
            'cuentos_dir_exists': cuentos_path.exists(),
            'canciones_dir_exists': canciones_path.exists(),
            'actividades_dir_exists': actividades_path.exists(),
            'cuentos_files': [f.name for f in cuentos_files],
            'canciones_files': [f.name for f in canciones_files],
            'actividades_files': [f.name for f in actividades_files],
            'total_files': len(cuentos_files) + len(canciones_files) + len(actividades_files)
        }
    }
    
    if rag_system:
        stats = rag_system.get_stats()
        status['vector_store'] = stats
    
    return status

# ============================================================================
# RUTAS PARA SERVIR EL FRONTEND
# ============================================================================

@app.get("/")
async def serve_index():
    """Servir index.html"""
    index_path = os.path.join(FRONTEND_DIR, "index.html")
    
    if os.path.exists(index_path):
        return FileResponse(index_path)
    else:
        return JSONResponse(
            content={
                "message": "ProfeGo API v2.0",
                "status": "running",
                "docs": "/docs",
                "error": "Frontend no encontrado"
            },
            status_code=404
        )

@app.get("/login.html")
async def serve_login():
    """Servir login.html"""
    login_path = os.path.join(FRONTEND_DIR, "login.html")
    
    if os.path.exists(login_path):
        return FileResponse(login_path)
    else:
        raise HTTPException(status_code=404, detail="login.html no encontrado")

@app.get("/menu.html")
async def serve_menu():
    """Servir menu.html"""
    menu_path = os.path.join(FRONTEND_DIR, "menu.html")
    
    if os.path.exists(menu_path):
        return FileResponse(menu_path)
    else:
        raise HTTPException(status_code=404, detail="menu.html no encontrado")

@app.get("/reset-password.html")
async def serve_reset_password():
    """Servir reset-password.html"""
    reset_path = os.path.join(FRONTEND_DIR, "reset-password.html")
    
    if os.path.exists(reset_path):
        return FileResponse(reset_path)
    else:
        raise HTTPException(status_code=404, detail="reset-password.html no encontrado")

@app.get("/forgot-password.html")
async def serve_forgot_password():
    """Servir forgot-password.html"""
    forgot_path = os.path.join(FRONTEND_DIR, "forgot-password.html")
    
    if os.path.exists(forgot_path):
        return FileResponse(forgot_path)
    else:
        raise HTTPException(status_code=404, detail="forgot-password.html no encontrado")

@app.get("/verification-success.html")
async def serve_verification_success():
    """Servir verification-success.html"""
    success_path = os.path.join(FRONTEND_DIR, "verification-success.html")
    
    if os.path.exists(success_path):
        return FileResponse(success_path)
    else:
        raise HTTPException(status_code=404, detail="verification-success.html no encontrado")

@app.get("/health")
async def health_check():
    """Verificar estado del servicio"""
    try:
        gcs_status = "connected" if gcs_storage.bucket.exists() else "disconnected"
        gemini_configured = bool(os.getenv("GEMINI_API_KEY"))
        
        # Verificar biblioteca RAG
        cuentos_count = len(list(Path('./rag_data/cuentos').glob('**/*.txt')))
        canciones_count = len(list(Path('./rag_data/canciones').glob('**/*.txt')))
        actividades_count = len(list(Path('./rag_data/actividades').glob('**/*.txt')))
        
        return {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "gcs_status": gcs_status,
            "bucket_name": gcs_storage.bucket_name,
            "frontend_dir": FRONTEND_DIR,
            "frontend_exists": os.path.exists(FRONTEND_DIR),
            "gemini_configured": gemini_configured,
            "rag_system": rag_system is not None,
            "rag_library": {
                "cuentos": cuentos_count,
                "canciones": canciones_count,
                "actividades": actividades_count,
                "total": cuentos_count + canciones_count + actividades_count
            },
            "version": "2.0.0"
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e)
        }

# ============================================================================
# PUNTO DE ENTRADA
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    
    print("=" * 60)
    print("ðŸš€ ProfeGo API v2.0 - Servidor Iniciando")
    print("=" * 60)
    print(f"ðŸ“ Frontend: {FRONTEND_DIR}")
    print(f"â˜ï¸ GCS Bucket: {gcs_storage.bucket_name}")
    print(f"ðŸ“¦ LÃ­mite de archivo: {MAX_FILE_SIZE / (1024*1024)}MB")
    print(f"ðŸ” CORS Origins: {allowed_origins}")
    print(f"ðŸ¤– Gemini AI: {'âœ… Configurado' if os.getenv('GEMINI_API_KEY') else 'âŒ No configurado'}")
    print(f"ðŸŒ Servidor: http://127.0.0.1:8000")
    print(f"ðŸ“– Docs: http://127.0.0.1:8000/docs")
    print("=" * 60)

    uvicorn.run(app, host="127.0.0.1", port=8000, log_level="info")
